from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT = r"D:\code\spider\docs\ts_ods_ods_news_fields.docx"


FIELDS = [
    ("1", "标准记录唯一标识", "record_id", "TEXT", "否", "主键", "标准记录唯一标识"),
    ("2", "数据来源", "data_source", "TEXT", "否", "", "采集源标识"),
    ("3", "数据类型", "data_type", "TEXT", "否", "固定值", "固定为 news"),
    ("4", "标题", "title", "TEXT", "是", "", "新闻标题"),
    ("5", "原文链接", "url", "TEXT", "是", "", "原文页面链接"),
    ("6", "源站链接", "source_url", "TEXT", "是", "", "源站或列表页链接"),
    ("7", "源发布时间", "source_published_at", "TIMESTAMPTZ", "是", "", "源数据发布时间"),
    ("8", "源更新时间", "source_updated_at", "TIMESTAMPTZ", "是", "", "源数据更新时间"),
    ("9", "摘要纯文本", "summary", "TEXT", "是", "", "标准化后的摘要文本"),
    ("10", "正文纯文本", "content", "TEXT", "是", "", "标准化后的正文文本"),
    ("11", "正文 HTML", "content_html", "TEXT", "是", "", "原始正文 HTML"),
    ("12", "摘要 HTML", "summary_html", "TEXT", "是", "", "原始摘要 HTML"),
    ("13", "作者", "author", "TEXT", "是", "", "作者信息"),
    (
        "14",
        "新闻类型",
        "news_type",
        "JSONB",
        "是",
        "",
        "新闻类型数组，目前仅 satellite_today 使用 category_names 写入",
    ),
    ("15", "机构信息", "organization", "JSONB", "是", "", "机构结构化信息"),
    ("16", "标签信息", "tags", "JSONB", "是", "", "标签数组"),
    ("17", "外链列表", "external_links", "JSONB", "是", "", "正文外部链接列表"),
    ("18", "附件列表", "attachments", "JSONB", "是", "", "附件资源列表"),
    ("19", "图片列表", "images", "JSONB", "是", "", "正文图片资源列表"),
    ("20", "轮播/组图列表", "slides", "JSONB", "是", "", "轮播图或组图资源列表"),
    ("21", "缩略图/封面图", "thumbnail", "TEXT", "是", "", "缩略图或封面图路径/链接"),
    ("22", "创建时间", "created_at", "TIMESTAMPTZ", "否", "默认 NOW()", "入库创建时间"),
    ("23", "更新时间", "updated_at", "TIMESTAMPTZ", "否", "默认 NOW()", "入库更新时间"),
]

HEADERS = ["序号", "名称", "字段", "类型", "是否可空", "约束", "备注"]
WIDTHS_CM = [1.2, 3.0, 4.0, 2.6, 2.2, 2.4, 6.0]


def set_run_font(run, size, bold=False):
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_cell_width(cell, width_cm):
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_cm / 2.54 * 1440)))


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run("数仓 ODS 表字段说明")
    set_run_font(title_run, 18, bold=True)

    name_p = doc.add_paragraph()
    name_run = name_p.add_run("表名：ts_ods.ods_news")
    set_run_font(name_run, 11, bold=True)

    desc_p = doc.add_paragraph()
    desc_run = desc_p.add_run("说明：ODS 层标准新闻表，用于承接标准化后的新闻数据。")
    set_run_font(desc_run, 10.5)

    table = doc.add_table(rows=1, cols=len(HEADERS))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for idx, header in enumerate(HEADERS):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_width(cell, WIDTHS_CM[idx])
        set_cell_fill(cell, "D9EAF7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, 10, bold=True)

    for row_data in FIELDS:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            cell.text = value
            set_cell_width(cell, WIDTHS_CM[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 3, 4) else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_run_font(run, 9.5)

    note = doc.add_paragraph()
    note_run = note.add_run(
        "注：news_type 为 JSONB 数组字段，当前仅 satellite_today 源写入，值来自 category_names。"
    )
    note_run.italic = True
    set_run_font(note_run, 9.5)

    doc.save(OUTPUT)


if __name__ == "__main__":
    main()
